from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("/Users/felipemarin/Desktop/ConnectRole/artifacts")
DOCX_PATH = OUTPUT_DIR / "connectrole-architecture-and-growth-plan.docx"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 45, 61)
MUTED = RGBColor(95, 107, 122)
LIGHT_FILL = "F4F6F9"
TABLE_HEAD = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
      node = tc_mar.find(qn(f"w:{tag}"))
      if node is None:
          node = OxmlElement(f"w:{tag}")
          tc_mar.append(node)
      node.set(qn("w:w"), str(value))
      node.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, color=DARK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Title", 24, DARK, 0, 6),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_header_footer(section, title_text):
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(title_text)
    set_font(run, size=9, color=MUTED)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    for cell in (left, right):
        set_cell_margins(cell, top=0, bottom=0)

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = p_left.add_run("ConnectRole architecture brief")
    set_font(run_left, size=9, color=MUTED)

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_right = p_right.add_run("Page ")
    set_font(run_right, size=9, color=MUTED)

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run_field = p_right.add_run()
    run_field._r.append(fld_char1)
    run_field._r.append(instr)
    run_field._r.append(fld_char2)
    set_font(run_field, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Architecture Brief")
    set_font(r, size=12, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title.add_run("ConnectRole")
    set_font(t1, size=28, color=DARK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    s1 = subtitle.add_run("Current system architecture, data flow, and a practical roadmap toward a production-ready cloud platform")
    set_font(s1, size=13, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    meta.columns[0].width = Inches(1.5)
    meta.columns[1].width = Inches(5.0)

    rows = [
        ("Prepared for", "Felipe Marin"),
        ("Project", "AI mock interview coach built with Next.js, TypeScript, LLM APIs, and browser voice tooling"),
        ("Purpose", "Summarize the current architecture and outline the next infrastructure investments that would strengthen both the product and the resume story"),
        ("Recommended direction", "PostgreSQL + Prisma + AWS-first deployment, with Azure as a strong secondary path"),
    ]
    for i, (label, value) in enumerate(rows):
        left = meta.rows[i].cells[0]
        right = meta.rows[i].cells[1]
        set_cell_shading(left, TABLE_HEAD)
        for cell in (left, right):
            set_cell_margins(cell)
        p1 = left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        set_font(r1, size=10.5, color=DARK, bold=True)
        p2 = right.paragraphs[0]
        r2 = p2.add_run(value)
        set_font(r2, size=10.5, color=DARK)

    doc.add_paragraph("")
    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(6.5)
    cell = callout.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "This document is written as a planning memo rather than a formal specification. The goal is to explain how ConnectRole works today, where the architecture is still prototype-oriented, and which backend and cloud additions would make the project feel much closer to something a company would run in production."
    )
    set_font(r, size=11, color=DARK)


def add_heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=11, color=DARK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_font(run, size=11, color=DARK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_font(run, size=11, color=DARK)


def add_two_col_table(doc, rows, widths=(2.0, 4.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])
    for i, (left_text, right_text) in enumerate(rows):
        left = table.rows[i].cells[0]
        right = table.rows[i].cells[1]
        set_cell_shading(left, TABLE_HEAD)
        for cell in (left, right):
            set_cell_margins(cell)
        p1 = left.paragraphs[0]
        r1 = p1.add_run(left_text)
        set_font(r1, size=10.5, color=DARK, bold=True)
        p2 = right.paragraphs[0]
        r2 = p2.add_run(right_text)
        set_font(r2, size=10.5, color=DARK)


def add_three_col_table(doc, headers, rows, widths=(1.5, 2.5, 2.5)):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_HEAD)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_font(r, size=10, color=DARK, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=10, color=DARK)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0], "ConnectRole | Architecture brief")
    add_cover(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(doc.sections[1], "ConnectRole | Current architecture")
    add_heading(doc, "1. What ConnectRole is today", 1)
    add_body(
        doc,
        "ConnectRole is currently a polished product prototype with a light backend. The user experience is strong: the app accepts a resume, parses a target job posting, generates interview questions, scores spoken answers, and produces a final report. Architecturally, though, the application is still closer to an advanced single-session workflow than a full production system."
    )
    add_body(
        doc,
        "The frontend is built in Next.js with the App Router, and most of the experience lives in three client-driven screens: setup, interview, and results. On the server side, API routes provide narrow services such as parsing setup inputs, generating opening questions, evaluating interview turns, summarizing the session, and calling Google Cloud Text-to-Speech. The intelligence layer lives in reusable modules under the lib directory, which keeps the project organized and makes the business logic easier to evolve."
    )
    add_heading(doc, "Current functional layers", 2)
    add_bullets(
        doc,
        [
            "Presentation layer: React components guide the user through resume upload, interview practice, and results review.",
            "Application layer: Next.js API routes receive structured requests from the browser and return parsed or generated data.",
            "Intelligence layer: helper modules handle resume parsing, job parsing, question generation, turn evaluation, and report creation.",
            "External services: the project already communicates with LLM providers and Google Cloud Text-to-Speech.",
        ],
    )
    add_heading(doc, "Why this is a good starting point", 2)
    add_body(
        doc,
        "From a hiring perspective, the existing app already shows real product thinking. It is not just a static portfolio site. It has a multi-step flow, API usage, typed data models, state transitions, and user-facing feedback logic. That gives it a strong base. The next leap is to turn the current transient workflow into a persistent, cloud-backed platform."
    )
    add_heading(doc, "Core files that define the system", 2)
    add_two_col_table(
        doc,
        [
            ("Setup flow", "components/setup-screen.tsx and app/api/setup/parse/route.ts"),
            ("Interview flow", "components/interview-screen.tsx and app/api/interview/turn/route.ts"),
            ("Results flow", "components/results-screen.tsx and app/api/interview/summary/route.ts"),
            ("Session storage", "lib/session.ts"),
            ("AI orchestration", "lib/setup-intelligence.ts, lib/interview-brain.ts, and lib/interview-engine.ts"),
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(doc.sections[2], "ConnectRole | Data flow")
    add_heading(doc, "2. How data moves through the system", 1)
    add_body(
        doc,
        "The current data flow is straightforward and easy to reason about. During setup, the browser extracts text from a PDF resume, gathers job-posting text from the user, and sends both payloads to a server route. The server parses and enriches those inputs, then returns structured resume and job objects. At that point, the browser becomes the main holder of state."
    )
    add_body(
        doc,
        "The interview page loads that saved setup data, asks questions, captures answers, and sends each turn to the evaluation route. The server evaluates the latest answer using the full client-supplied context, then returns the next question and coaching feedback. Finally, the results page assembles a final report, asks the server for a summary, and stores the completed report in the browser as well."
    )
    add_heading(doc, "Current state ownership", 2)
    add_three_col_table(
        doc,
        ["Stage", "Where data is created", "Where data lives after creation"],
        [
            ("Setup", "Browser + /api/setup/parse", "sessionStorage in the browser"),
            ("Interview turns", "Browser + /api/interview/turn", "sessionStorage in the browser"),
            ("Results", "Browser + /api/interview/summary", "sessionStorage in the browser"),
            ("Voice output", "/api/tts", "Returned directly as audio bytes"),
        ],
    )
    add_heading(doc, "What this means in practice", 2)
    add_bullets(
        doc,
        [
            "The app behaves well for one session on one device.",
            "There is no durable interview history if the tab is lost or the user changes devices.",
            "There is no server-side source of truth, so analytics and user accounts would be hard to add cleanly later.",
            "The client sends large context payloads repeatedly, which is acceptable for a prototype but not ideal for scale.",
        ],
    )
    add_heading(doc, "Most important architectural gap", 2)
    add_body(
        doc,
        "The largest gap is not the lack of a cloud provider by itself. It is the lack of persistence and ownership boundaries. Right now, the browser effectively owns interview state. In a production system, the server should own interview sessions and the database should preserve them. Once that shift happens, cloud infrastructure starts to make much more sense because there is a real backend to host, monitor, protect, and scale."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(doc.sections[3], "ConnectRole | Target architecture")
    add_heading(doc, "3. Recommended target architecture", 1)
    add_body(
        doc,
        "The best near-term architecture for ConnectRole is a server-centered design with PostgreSQL as the durable store, object storage for uploaded files, and managed cloud services for deployment and observability. This does not require a giant rewrite. It is more of a shift in responsibility: moving long-lived state out of sessionStorage and into backend services."
    )
    add_heading(doc, "Recommended stack", 2)
    add_two_col_table(
        doc,
        [
            ("Web application", "Next.js App Router, keeping the existing frontend and API route structure"),
            ("Database", "PostgreSQL with Prisma for schema management and typed queries"),
            ("File storage", "Amazon S3 for resume uploads and future exported reports"),
            ("Authentication", "Auth.js or Amazon Cognito, depending on how cloud-native you want the story to be"),
            ("Deployment", "AWS Amplify or a container-based deployment path"),
            ("Monitoring", "CloudWatch, request logging, and basic error tracking"),
        ],
    )
    add_heading(doc, "How the improved data model should work", 2)
    add_numbered(
        doc,
        [
            "A user signs in and starts a new interview preparation session.",
            "The resume PDF is uploaded to object storage and its extracted text is saved in PostgreSQL.",
            "The parsed job posting is stored as a first-class database record rather than only a temporary payload.",
            "Starting an interview creates an interview_session record on the server.",
            "Each answer inserts a new interview_turn row and updates the current server-side session state.",
            "Completing the interview creates a final_report row that can be revisited later.",
        ],
    )
    add_heading(doc, "Why PostgreSQL is the right database here", 2)
    add_body(
        doc,
        "This project is fundamentally relational. Users own resumes. Resumes and job postings connect to interview sessions. Sessions contain many turns. Sessions also produce one final report. That shape fits PostgreSQL naturally. It is also one of the most common databases employers expect students and new grads to understand, especially when paired with Prisma, SQL, and cloud deployment."
    )
    add_heading(doc, "AWS versus Azure", 2)
    add_three_col_table(
        doc,
        ["Concern", "AWS-first path", "Azure-first path"],
        [
            ("Hosting", "Amplify or container deployment", "App Service or Static Web Apps"),
            ("Database", "RDS for PostgreSQL", "Azure Database for PostgreSQL"),
            ("Object storage", "S3", "Blob Storage"),
            ("Identity", "Cognito", "Microsoft Entra"),
            ("Monitoring", "CloudWatch", "Application Insights / Azure Monitor"),
        ],
    )
    add_body(
        doc,
        "If only one path is chosen, AWS is the stronger general-purpose choice for this project because it aligns well with full-stack deployment, object storage, and the broader job market. Azure remains a good secondary path, especially for companies built around Microsoft tooling."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(doc.sections[4], "ConnectRole | Improvement plan")
    add_heading(doc, "4. Improvement roadmap for the next few months", 1)
    add_body(
        doc,
        "A good improvement plan should create visible momentum quickly. The goal is not to boil the ocean. It is to turn the project into something that demonstrates backend ownership, database design, deployment maturity, and real cloud familiarity in a sequence that still feels achievable for a senior student."
    )
    add_heading(doc, "Phase-by-phase plan", 2)
    add_three_col_table(
        doc,
        ["Phase", "Main deliverable", "Why it matters"],
        [
            ("1", "Prisma + PostgreSQL integration", "Introduces durable backend state and shows database modeling skill"),
            ("2", "User authentication", "Turns anonymous sessions into user-owned history"),
            ("3", "Resume file storage in S3", "Adds real cloud storage and separation of file data from relational data"),
            ("4", "Docker + local compose setup", "Shows environment reproducibility and deployment readiness"),
            ("5", "AWS deployment + monitoring", "Makes the project production-like and resume-ready"),
        ],
    )
    add_heading(doc, "The highest-value implementation order", 2)
    add_bullets(
        doc,
        [
            "Move setup, interview session, and final report persistence from sessionStorage to PostgreSQL.",
            "Introduce an interview session identifier so the browser can reload state from the server.",
            "Store the original PDF file separately from the parsed text and metadata.",
            "Containerize the application so local development and deployment use the same runtime assumptions.",
            "Add a basic CI pipeline that runs linting, type checks, and build validation.",
        ],
    )
    add_heading(doc, "How this helps your resume", 2)
    add_body(
        doc,
        "Employers care less about whether a student merely used buzzwords and more about whether the system design choices make sense. If ConnectRole evolves into a project with PostgreSQL, Docker, object storage, authentication, and a managed AWS deployment, you can talk concretely about tradeoffs: why some data belongs in relational tables, why files belong in object storage, why Docker helps standardize environments, and how cloud services reduce operational burden."
    )
    add_body(
        doc,
        "That is exactly the kind of conversation that tends to differentiate a candidate who has built projects from a candidate who has only taken classes. The strongest version of this story is not just 'I used AWS.' It is 'I redesigned a client-held prototype into a persistent cloud-backed system and can explain each infrastructure choice.'"
    )
    add_heading(doc, "Final recommendation", 2)
    add_body(
        doc,
        "The clearest next move is to make ConnectRole a database-backed AWS project. Start with PostgreSQL and Prisma, then add auth, S3, Docker, and deployment. That sequence improves the product, makes the architecture more correct, and gives you several highly marketable skills at once without losing the momentum of the app you already built."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
